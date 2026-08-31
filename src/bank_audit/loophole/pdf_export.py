"""Экспорт результата loophole в PDF через Playwright.

Переиспользует стиль web/pdf_export (Source Serif 4 / Geist / JetBrains Mono).
Генерирует HTML из записей, рендерит в A4 PDF через headless Chromium.
"""
from __future__ import annotations

import logging
from html import escape
from io import BytesIO

log = logging.getLogger(__name__)


_HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="ru"><head><meta charset="utf-8">
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="stylesheet" href="https://fonts.googleapis.com/css2?family=Source+Serif+4:opsz,wght@8..60,400;8..60,600&family=Geist:wght@400;500;600&family=JetBrains+Mono:wght@400;500&display=swap">
<style>
  body {{ font-family: "Geist", system-ui, sans-serif; color: #1a1a1a; max-width: 720px; margin: 0 auto; padding: 32px; }}
  h1 {{ font-family: "Source Serif 4", Georgia, serif; font-size: 1.5rem; }}
  h2 {{ font-family: "Source Serif 4", Georgia, serif; font-size: 1.1rem; margin-top: 18px; }}
  .meta {{ color: #6b6b6b; font-size: 0.85rem; margin-bottom: 20px; }}
  .record {{ border-bottom: 1px solid #d8d8d2; padding: 12px 0; }}
  .record .title {{ font-weight: 600; }}
  .record .url {{ font-family: "JetBrains Mono", monospace; font-size: 0.8rem; color: #6b6b6b; word-break: break-all; }}
  .record .verdict {{ margin-top: 4px; }}
  .loophole {{ color: #b03a2e; }}
</style></head><body>
<h1>Отчёт: лазейки и уязвимости</h1>
<div class="meta">Сформирован: {generated_at}</div>
{records_html}
</body></html>"""


def _record_html(r: dict) -> str:
    title = r.get("title") or r.get("snippet") or "(без названия)"
    url = r.get("url") or ""
    bank = r.get("bank_slug") or "—"
    is_l = r.get("is_loophole")
    verdict_cls = "loophole" if is_l else ""
    verdict = "лазейка" if is_l else "не лазейка"
    conf = r.get("verdict_confidence")
    conf_str = f" (доверие {conf:.2f})" if conf is not None else ""
    return (
        f'<div class="record">'
        f'<div class="title">{title}</div>'
        f'<div class="url">{url}</div>'
        f'<div>Банк: {bank}</div>'
        f'<div class="verdict {verdict_cls}">Вердикт: {verdict}{conf_str}</div>'
        f'</div>'
    )


def render_html(records: list[dict], *, generated_at: str = "") -> str:
    """Возвращает HTML отчёта."""
    from ..clock import today_ru
    gen = generated_at or today_ru()
    records_html = "\n".join(_record_html(r) for r in records)
    return _HTML_TEMPLATE.format(generated_at=gen, records_html=records_html)


async def export_pdf(records: list[dict], *, output_path: str = "") -> bytes:
    """Рендерит HTML в PDF через Playwright. Возвращает PDF-байты.

    Если Playwright недоступен — падает с понятной ошибкой.
    """
    html = render_html(records)
    try:
        from playwright.async_api import async_playwright
    except Exception as e:
        raise RuntimeError(f"Playwright недоступен: {e}") from e
    async with async_playwright() as p:
        browser = await p.chromium.launch()
        page = await browser.new_page()
        await page.set_content(html, wait_until="networkidle")
        pdf = await page.pdf(format="A4", print_background=True)
        await browser.close()
    if output_path:
        from pathlib import Path
        Path(output_path).write_bytes(pdf)
    return pdf


def render_research_report_html(report: dict) -> str:
    """Собирает отдельный безопасный HTML для immutable отчёта исследования."""
    def paragraph(value: object) -> str:
        return "<br>".join(escape(line) for line in str(value or "").splitlines()) or "—"

    evidence = report.get("evidence") or []
    if evidence:
        evidence_html = "".join(
            "<li><strong>{title}</strong><br><span class=\"url\">{url}</span>"
            "<p>{text}</p></li>".format(
                title=paragraph(item.get("title") or "Источник без названия"),
                url=paragraph(item.get("url")),
                text=paragraph(item.get("extracted_text")),
            )
            for item in evidence
            if isinstance(item, dict)
        ) or "<p>Проверенные доказательства отсутствуют.</p>"
    else:
        evidence_html = "<p>Проверенные доказательства отсутствуют.</p>"
    return """<!doctype html><html lang=\"ru\"><head><meta charset=\"utf-8\"><style>
body {{ font-family: Arial, sans-serif; color: #1a1a1a; max-width: 720px; margin: 0 auto; padding: 32px; }}
h1 {{ font-size: 24px; }} h2 {{ margin-top: 24px; }} .url {{ word-break: break-all; color: #555; }}
</style></head><body><h1>Отчёт AI-исследования</h1><h2>Тема</h2><p>{query}</p>
<h2>Итог</h2><p>{result}</p><h2>Проверенные доказательства и источники</h2><ul>{evidence}</ul>
</body></html>""".format(
        query=paragraph(report.get("query")), result=paragraph(report.get("result")), evidence=evidence_html
    )


async def export_research_report_pdf(report: dict) -> bytes:
    """Рендерит PDF отчёта исследования, не используя небезопасный catalog renderer."""
    try:
        from playwright.async_api import async_playwright
    except Exception as exc:
        raise RuntimeError(f"Playwright недоступен: {exc}") from exc
    async with async_playwright() as playwright:
        browser = await playwright.chromium.launch()
        try:
            page = await browser.new_page()
            await page.set_content(render_research_report_html(report), wait_until="networkidle")
            return await page.pdf(format="A4", print_background=True)
        finally:
            await browser.close()


def export_research_report_docx(report: dict) -> bytes:
    """Создаёт Word-документ из тех же immutable данных, что и PDF."""
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError(f"Word-экспорт недоступен: {exc}") from exc
    document = Document()
    document.add_heading("Отчёт AI-исследования", level=0)
    document.add_heading("Тема", level=1)
    document.add_paragraph(str(report.get("query") or ""))
    document.add_heading("Итог", level=1)
    document.add_paragraph(str(report.get("result") or ""))
    document.add_heading("Проверенные доказательства и источники", level=1)
    evidence = report.get("evidence") or []
    if not evidence:
        document.add_paragraph("Проверенные доказательства отсутствуют.")
    for item in evidence:
        if not isinstance(item, dict):
            continue
        document.add_paragraph(str(item.get("title") or "Источник без названия"), style="List Bullet")
        document.add_paragraph(str(item.get("url") or ""))
        if item.get("extracted_text"):
            document.add_paragraph(str(item["extracted_text"]))
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
